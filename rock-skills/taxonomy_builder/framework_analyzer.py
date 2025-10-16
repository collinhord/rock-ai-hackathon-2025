"""
Framework Analyzer Module

Parses academic frameworks and educational documents, extracts taxonomies,
and compares them against our taxonomy.

Supports: PDF, DOCX, TXT formats
"""

from pathlib import Path
from typing import Dict, List, Optional, Union
from dataclasses import dataclass
import json

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not available. PDF parsing will be disabled.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX parsing will be disabled.")

from llm_interface import LLMInterface
from compatibility import TaxonomyAccess


@dataclass
class FrameworkDocument:
    """Represents a parsed framework document."""
    filename: str
    content: str
    format: str
    word_count: int
    metadata: Dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class FrameworkParser:
    """Parses framework documents in various formats."""
    
    @staticmethod
    def parse_file(file_path: Union[str, Path]) -> FrameworkDocument:
        """
        Parse a framework file.
        
        Args:
            file_path: Path to framework file
            
        Returns:
            FrameworkDocument with parsed content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return FrameworkParser._parse_pdf(file_path)
        elif suffix == '.docx':
            return FrameworkParser._parse_docx(file_path)
        elif suffix in ['.txt', '.md']:
            return FrameworkParser._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    @staticmethod
    def _parse_pdf(file_path: Path) -> FrameworkDocument:
        """Parse PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is required to parse PDF files. Install with: pip install PyPDF2")
        
        content = []
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            metadata = {
                'page_count': len(reader.pages),
                'pdf_metadata': reader.metadata if reader.metadata else {}
            }
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content.append(text)
        
        full_content = '\n\n'.join(content)
        
        return FrameworkDocument(
            filename=file_path.name,
            content=full_content,
            format='pdf',
            word_count=len(full_content.split()),
            metadata=metadata
        )
    
    @staticmethod
    def _parse_docx(file_path: Path) -> FrameworkDocument:
        """Parse DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required to parse DOCX files. Install with: pip install python-docx")
        
        doc = Document(file_path)
        
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                if row_text.strip():
                    content.append(row_text)
        
        full_content = '\n\n'.join(content)
        
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables)
        }
        
        return FrameworkDocument(
            filename=file_path.name,
            content=full_content,
            format='docx',
            word_count=len(full_content.split()),
            metadata=metadata
        )
    
    @staticmethod
    def _parse_text(file_path: Path) -> FrameworkDocument:
        """Parse text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return FrameworkDocument(
            filename=file_path.name,
            content=content,
            format='text',
            word_count=len(content.split()),
            metadata={}
        )


class FrameworkAnalyzer:
    """Analyzes frameworks and compares against our taxonomy."""
    
    def __init__(self, taxonomy_access: TaxonomyAccess, llm: LLMInterface):
        """
        Initialize analyzer.
        
        Args:
            taxonomy_access: TaxonomyAccess instance
            llm: LLMInterface instance
        """
        self.tax = taxonomy_access
        self.llm = llm
    
    def analyze_framework(self, 
                         file_path: Union[str, Path],
                         compare_with_taxonomy: bool = True) -> Dict:
        """
        Analyze a framework document.
        
        Args:
            file_path: Path to framework file
            compare_with_taxonomy: Whether to compare with our taxonomy
            
        Returns:
            Dictionary with analysis results
        """
        print(f"Parsing framework: {file_path}")
        
        # Parse document
        doc = FrameworkParser.parse_file(file_path)
        
        print(f"  Format: {doc.format}")
        print(f"  Word count: {doc.word_count}")
        
        # Extract taxonomy from framework
        print("\nExtracting taxonomy structure from framework...")
        extracted_taxonomy = self.llm.extract_framework(doc.content)
        
        if compare_with_taxonomy:
            print("\nComparing with our taxonomy...")
            comparison = self._compare_with_our_taxonomy(extracted_taxonomy)
        else:
            comparison = None
        
        return {
            'document': {
                'filename': doc.filename,
                'format': doc.format,
                'word_count': doc.word_count,
                'metadata': doc.metadata
            },
            'extracted_taxonomy': extracted_taxonomy,
            'comparison': comparison,
            'llm_usage': self.llm.get_usage_stats()
        }
    
    def _compare_with_our_taxonomy(self, extracted_taxonomy: Dict) -> Dict:
        """
        Compare extracted framework taxonomy with ours.
        
        Args:
            extracted_taxonomy: Taxonomy extracted from framework
            
        Returns:
            Comparison results
        """
        # Get summary of our taxonomy
        our_summary = self._get_taxonomy_summary()
        
        # Format extracted taxonomy for comparison
        framework_summary = json.dumps(extracted_taxonomy, indent=2)
        
        # Use LLM to compare
        comparison = self.llm.compare_taxonomies(our_summary, framework_summary)
        
        return comparison
    
    def _get_taxonomy_summary(self) -> str:
        """Get formatted summary of our taxonomy."""
        summary_parts = []
        
        # Get all strands
        strands = self.tax.get_nodes_by_level('Strand')
        
        summary_parts.append("# Our Taxonomy Structure\n")
        
        for strand in strands:
            summary_parts.append(f"\n## {strand.name}")
            
            # Get pillars under this strand
            pillars = self.tax.get_children(strand.uuid)
            
            for pillar in pillars:
                summary_parts.append(f"\n### {pillar.name}")
                
                # Get domains under this pillar
                domains = self.tax.get_children(pillar.uuid)
                
                for domain in domains[:5]:  # Limit to first 5 to keep summary manageable
                    summary_parts.append(f"  - {domain.name}")
                
                if len(domains) > 5:
                    summary_parts.append(f"  - ... and {len(domains) - 5} more")
        
        return '\n'.join(summary_parts)
    
    def batch_analyze(self, file_paths: List[Union[str, Path]]) -> List[Dict]:
        """
        Analyze multiple framework files.
        
        Args:
            file_paths: List of paths to framework files
            
        Returns:
            List of analysis results
        """
        results = []
        
        for i, file_path in enumerate(file_paths, 1):
            print(f"\n{'='*60}")
            print(f"Analyzing framework {i}/{len(file_paths)}")
            print(f"{'='*60}")
            
            try:
                result = self.analyze_framework(file_path)
                results.append(result)
            except Exception as e:
                print(f"Error analyzing {file_path}: {e}")
                results.append({
                    'error': str(e),
                    'file_path': str(file_path)
                })
        
        return results
    
    def generate_report(self, analysis_results: Union[Dict, List[Dict]], 
                       output_path: Optional[Path] = None) -> str:
        """
        Generate markdown report from analysis results.
        
        Args:
            analysis_results: Single or multiple analysis results
            output_path: Optional path to save report
            
        Returns:
            Markdown report string
        """
        if isinstance(analysis_results, dict):
            analysis_results = [analysis_results]
        
        lines = []
        lines.append("# Framework Analysis Report")
        lines.append("")
        lines.append(f"**Total Frameworks Analyzed:** {len(analysis_results)}")
        lines.append("")
        
        for i, result in enumerate(analysis_results, 1):
            if 'error' in result:
                lines.append(f"## Framework {i}: ERROR")
                lines.append(f"**File:** {result.get('file_path', 'unknown')}")
                lines.append(f"**Error:** {result['error']}")
                lines.append("")
                continue
            
            doc = result['document']
            lines.append(f"## Framework {i}: {doc['filename']}")
            lines.append("")
            lines.append(f"**Format:** {doc['format']}")
            lines.append(f"**Word Count:** {doc['word_count']}")
            lines.append("")
            
            # Extracted taxonomy
            lines.append("### Extracted Taxonomy")
            lines.append("```json")
            lines.append(json.dumps(result['extracted_taxonomy'], indent=2))
            lines.append("```")
            lines.append("")
            
            # Comparison if available
            if result.get('comparison'):
                lines.append("### Comparison with Our Taxonomy")
                lines.append("```json")
                lines.append(json.dumps(result['comparison'], indent=2))
                lines.append("```")
                lines.append("")
        
        # LLM usage summary
        if analysis_results and 'llm_usage' in analysis_results[0]:
            lines.append("## LLM Usage Summary")
            lines.append("")
            usage = analysis_results[0]['llm_usage']
            lines.append(f"- Total tokens: {usage['total_tokens']}")
            lines.append(f"- Estimated cost: ${usage['estimated_cost_usd']}")
            lines.append("")
        
        report = '\n'.join(lines)
        
        if output_path:
            with open(output_path, 'w') as f:
                f.write(report)
            print(f"\nReport saved to: {output_path}")
        
        return report


# Example usage
if __name__ == '__main__':
    print("=== Framework Analyzer Demo ===\n")
    print("This module requires framework documents to analyze.")
    print("\nUsage:")
    print("  from framework_analyzer import FrameworkAnalyzer")
    print("  from compatibility import TaxonomyAccess")
    print("  from llm_interface import LLMInterface")
    print("")
    print("  with TaxonomyAccess() as tax:")
    print("      llm = LLMInterface(provider='bedrock')")
    print("      analyzer = FrameworkAnalyzer(tax, llm)")
    print("      result = analyzer.analyze_framework('path/to/framework.pdf')")
    print("      report = analyzer.generate_report(result, 'framework_report.md')")

