"""
Framework Analyzer Module

Parses academic frameworks and educational documents, extracts taxonomies,
and compares them against our taxonomy.

Supports: PDF, DOCX, TXT formats
"""

from pathlib import Path
from typing import Dict, List, Optional, Union
from dataclasses import dataclass
import json

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not available. PDF parsing will be disabled.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX parsing will be disabled.")

from llm_interface import LLMInterface
from compatibility import TaxonomyAccess


@dataclass
class FrameworkDocument:
    """Represents a parsed framework document."""
    filename: str
    content: str
    format: str
    word_count: int
    metadata: Dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class FrameworkParser:
    """Parses framework documents in various formats."""
    
    @staticmethod
    def parse_file(file_path: Union[str, Path]) -> FrameworkDocument:
        """
        Parse a framework file.
        
        Args:
            file_path: Path to framework file
            
        Returns:
            FrameworkDocument with parsed content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return FrameworkParser._parse_pdf(file_path)
        elif suffix == '.docx':
            return FrameworkParser._parse_docx(file_path)
        elif suffix in ['.txt', '.md']:
            return FrameworkParser._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    @staticmethod
    def _parse_pdf(file_path: Path) -> FrameworkDocument:
        """Parse PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is required to parse PDF files. Install with: pip install PyPDF2")
        
        content = []
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            metadata = {
                'page_count': len(reader.pages),
                'pdf_metadata': reader.metadata if reader.metadata else {}
            }
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content.append(text)
        
        full_content = '\n\n'.join(content)
        
        return FrameworkDocument(
            filename=file_path.name,
            content=full_content,
            format='pdf',
            word_count=len(full_content.split()),
            metadata=metadata
        )
    
    @staticmethod
    def _parse_docx(file_path: Path) -> FrameworkDocument:
        """Parse DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required to parse DOCX files. Install with: pip install python-docx")
        
        doc = Document(file_path)
        
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                if row_text.strip():
                    content.append(row_text)
        
        full_content = '\n\n'.join(content)
        
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables)
        }
        
        return FrameworkDocument(
            filename=file_path.name,
            content=full_content,
            format='docx',
            word_count=len(full_content.split()),
            metadata=metadata
        )
    
    @staticmethod
    def _parse_text(file_path: Path) -> FrameworkDocument:
        """Parse text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return FrameworkDocument(
            filename=file_path.name,
            content=content,
            format='text',
            word_count=len(content.split()),
            metadata={}
        )


class FrameworkAnalyzer:
    """Analyzes frameworks and compares against our taxonomy."""
    
    def __init__(self, taxonomy_access: TaxonomyAccess = None, llm: LLMInterface = None):
        """
        Initialize analyzer.
        
        Args:
            taxonomy_access: TaxonomyAccess instance (optional for new taxonomy creation)
            llm: LLMInterface instance (optional, will create if not provided)
        """
        self.tax = taxonomy_access
        if llm is None:
            # Create default LLM interface
            self.llm = LLMInterface(provider='bedrock')
        else:
            self.llm = llm
    
    def analyze_framework(self, 
                         file_path: Union[str, Path],
                         compare_with_taxonomy: bool = True) -> Dict:
        """
        Analyze a framework document.
        
        Args:
            file_path: Path to framework file
            compare_with_taxonomy: Whether to compare with our taxonomy
            
        Returns:
            Dictionary with analysis results
        """
        print(f"Parsing framework: {file_path}")
        
        # Parse document
        doc = FrameworkParser.parse_file(file_path)
        
        print(f"  Format: {doc.format}")
        print(f"  Word count: {doc.word_count}")
        
        # Extract taxonomy from framework
        print("\nExtracting taxonomy structure from framework...")
        extracted_taxonomy = self.llm.extract_framework(doc.content)
        
        if compare_with_taxonomy:
            print("\nComparing with our taxonomy...")
            comparison = self._compare_with_our_taxonomy(extracted_taxonomy)
        else:
            comparison = None
        
        return {
            'document': {
                'filename': doc.filename,
                'format': doc.format,
                'word_count': doc.word_count,
                'metadata': doc.metadata
            },
            'extracted_taxonomy': extracted_taxonomy,
            'comparison': comparison,
            'llm_usage': self.llm.get_usage_stats()
        }
    
    def _compare_with_our_taxonomy(self, extracted_taxonomy: Dict) -> Dict:
        """
        Compare extracted framework taxonomy with ours.
        
        Args:
            extracted_taxonomy: Taxonomy extracted from framework
            
        Returns:
            Comparison results
        """
        # Get summary of our taxonomy
        our_summary = self._get_taxonomy_summary()
        
        # Format extracted taxonomy for comparison
        framework_summary = json.dumps(extracted_taxonomy, indent=2)
        
        # Use LLM to compare
        comparison = self.llm.compare_taxonomies(our_summary, framework_summary)
        
        return comparison
    
    def _get_taxonomy_summary(self) -> str:
        """Get formatted summary of our taxonomy."""
        summary_parts = []
        
        # Get all strands
        strands = self.tax.get_nodes_by_level('Strand')
        
        summary_parts.append("# Our Taxonomy Structure\n")
        
        for strand in strands:
            summary_parts.append(f"\n## {strand.name}")
            
            # Get pillars under this strand
            pillars = self.tax.get_children(strand.uuid)
            
            for pillar in pillars:
                summary_parts.append(f"\n### {pillar.name}")
                
                # Get domains under this pillar
                domains = self.tax.get_children(pillar.uuid)
                
                for domain in domains[:5]:  # Limit to first 5 to keep summary manageable
                    summary_parts.append(f"  - {domain.name}")
                
                if len(domains) > 5:
                    summary_parts.append(f"  - ... and {len(domains) - 5} more")
        
        return '\n'.join(summary_parts)
    
    def batch_analyze(self, file_paths: List[Union[str, Path]]) -> List[Dict]:
        """
        Analyze multiple framework files.
        
        Args:
            file_paths: List of paths to framework files
            
        Returns:
            List of analysis results
        """
        results = []
        
        for i, file_path in enumerate(file_paths, 1):
            print(f"\n{'='*60}")
            print(f"Analyzing framework {i}/{len(file_paths)}")
            print(f"{'='*60}")
            
            try:
                result = self.analyze_framework(file_path)
                results.append(result)
            except Exception as e:
                print(f"Error analyzing {file_path}: {e}")
                results.append({
                    'error': str(e),
                    'file_path': str(file_path)
                })
        
        return results
    
    def extract_adaptive_taxonomy(self, 
                                 file_path: Union[str, Path],
                                 subject_area: str = 'ela') -> Dict:
        """
        Extract taxonomy with adaptive hierarchy detection.
        
        Args:
            file_path: Path to framework document
            subject_area: Subject area ('ela', 'math', 'science', 'general')
            
        Returns:
            Dictionary with extracted taxonomy structure
        """
        print(f"\n{'='*60}")
        print(f"ADAPTIVE TAXONOMY EXTRACTION")
        print(f"{'='*60}")
        print(f"File: {file_path}")
        print(f"Subject: {subject_area}")
        
        # Parse document
        doc = FrameworkParser.parse_file(file_path)
        
        print(f"\nDocument Info:")
        print(f"  Format: {doc.format}")
        print(f"  Word count: {doc.word_count:,}")
        
        # Extract taxonomy using LLM
        print(f"\nExtracting taxonomy structure (this may take 1-2 minutes)...")
        extracted = self.llm.extract_taxonomy_adaptive(doc.content)
        
        if extracted.get('parse_error'):
            print("\n⚠️  Warning: LLM response could not be parsed as JSON")
            print("See raw_response in output for details")
        else:
            print("\n✅ Taxonomy extraction complete!")
            if 'metadata_summary' in extracted:
                summary = extracted['metadata_summary']
                print(f"\nExtracted:")
                print(f"  - {summary.get('total_concepts_extracted', 0)} concepts")
                if 'concepts_by_level' in summary:
                    print(f"  - Hierarchy levels: {', '.join(summary['concepts_by_level'].keys())}")
        
        return {
            'document': {
                'filename': doc.filename,
                'format': doc.format,
                'word_count': doc.word_count,
                'subject_area': subject_area
            },
            'extraction': extracted
        }
    
    def generate_master_concepts(self, 
                                extracted_taxonomy: Dict,
                                subject_area: str = 'ela',
                                id_prefix: str = None) -> List[Dict]:
        """
        Convert extracted taxonomy to master concepts format.
        
        Args:
            extracted_taxonomy: Output from extract_adaptive_taxonomy
            subject_area: Subject area code ('ela', 'math', 'science')
            id_prefix: Custom ID prefix (default: MC-{SUBJECT}-
            
        Returns:
            List of master concept dictionaries matching master-concepts.csv format
        """
        print(f"\n{'='*60}")
        print(f"GENERATING MASTER CONCEPTS")
        print(f"{'='*60}")
        
        if 'extraction' in extracted_taxonomy:
            extraction = extracted_taxonomy['extraction']
        else:
            extraction = extracted_taxonomy
        
        if extraction.get('parse_error'):
            raise ValueError("Cannot generate master concepts from failed extraction")
        
        concepts = extraction.get('extracted_concepts', [])
        proposed_hierarchy = extraction.get('proposed_hierarchy', {})
        
        print(f"Processing {len(concepts)} extracted concepts...")
        
        # Generate ID prefix
        if id_prefix is None:
            subject_code = subject_area.upper()
            id_prefix = f"MC-{subject_code}-"
        
        master_concepts = []
        
        for i, concept in enumerate(concepts, start=1):
            # Generate ID
            concept_id = f"{id_prefix}{i:04d}"
            
            # Map to master concept fields
            master_concept = {
                'MASTER_CONCEPT_ID': concept_id,
                'MASTER_CONCEPT_NAME': concept.get('name', ''),
                'DESCRIPTION': concept.get('description', ''),
                'COMPLEXITY_BAND': concept.get('complexity_band', ''),
                'SKILL_COUNT': 0,  # Will be populated during mapping
                'AUTHORITY_COUNT': 0,  # Will be populated during mapping
                'GRADE_RANGE': concept.get('grade_range', ''),
                'TAXONOMY_CONFIDENCE': 'Medium',  # Default, can be refined
            }
            
            # Map hierarchy levels to SOR fields if applicable
            level = concept.get('level', '')
            parent = concept.get('parent_concept', '')
            
            # Try to map to Strand/Pillar/Domain structure
            # This is heuristic and subject-specific
            if subject_area == 'ela':
                # ELA-specific mapping logic
                if level in ['Strand', 'Level 1']:
                    master_concept['SOR_STRAND'] = concept.get('name', '')
                    master_concept['SOR_PILLAR'] = ''
                    master_concept['SOR_DOMAIN'] = ''
                elif level in ['Sub-strand', 'Pillar', 'Level 2']:
                    master_concept['SOR_STRAND'] = parent
                    master_concept['SOR_PILLAR'] = concept.get('name', '')
                    master_concept['SOR_DOMAIN'] = ''
                elif level in ['Domain', 'Concept', 'Level 3']:
                    # Try to infer from parent chain
                    master_concept['SOR_STRAND'] = ''
                    master_concept['SOR_PILLAR'] = parent
                    master_concept['SOR_DOMAIN'] = concept.get('name', '')
                else:
                    master_concept['SOR_STRAND'] = ''
                    master_concept['SOR_PILLAR'] = ''
                    master_concept['SOR_DOMAIN'] = concept.get('name', '')
            else:
                # For non-ELA, leave SOR fields empty or use generic mapping
                master_concept['SOR_STRAND'] = ''
                master_concept['SOR_PILLAR'] = ''
                master_concept['SOR_DOMAIN'] = concept.get('name', '')
            
            # Extract subject-specific metadata
            subject_metadata = concept.get('subject_metadata', {})
            
            if subject_area == 'ela':
                master_concept['TEXT_TYPE'] = subject_metadata.get('text_type', '')
                master_concept['TEXT_MODE'] = subject_metadata.get('text_mode', '')
                master_concept['SKILL_DOMAIN'] = subject_metadata.get('skill_domain', 'reading')
            elif subject_area == 'math':
                master_concept['MATHEMATICAL_DOMAIN'] = subject_metadata.get('mathematical_domain', '')
                master_concept['COGNITIVE_DEMAND'] = subject_metadata.get('cognitive_demand', '')
                master_concept['REPRESENTATION'] = subject_metadata.get('representation', '')
            
            # Add prerequisite and equivalence tracking
            prerequisites = concept.get('prerequisites', [])
            if prerequisites:
                master_concept['PREREQUISITE_CONCEPT_ID'] = ','.join([f"{id_prefix}{j:04d}" for j in range(1, len(prerequisites)+1)])
            else:
                master_concept['PREREQUISITE_CONCEPT_ID'] = ''
            
            master_concept['EQUIVALENCE_GROUP_ID'] = ''  # To be determined during alignment
            
            master_concepts.append(master_concept)
        
        print(f"✅ Generated {len(master_concepts)} master concepts")
        return master_concepts
    
    def analyze_alignment(self, 
                         extracted_concepts: List[Dict],
                         existing_taxonomy_df: Optional[object] = None) -> Dict:
        """
        Analyze alignment between extracted concepts and existing taxonomy.
        
        Args:
            extracted_concepts: List of master concepts from generate_master_concepts
            existing_taxonomy_df: DataFrame of existing taxonomy (or None for new taxonomy)
            
        Returns:
            Dictionary with alignment analysis
        """
        print(f"\n{'='*60}")
        print(f"ALIGNMENT ANALYSIS")
        print(f"{'='*60}")
        
        if existing_taxonomy_df is None and self.tax is not None:
            existing_taxonomy_df = self.tax.get_taxonomy_df()
        
        analysis = {
            'extracted_concept_count': len(extracted_concepts),
            'existing_concept_count': 0,
            'alignment_score': 0.0,
            'missing_from_taxonomy': [],
            'missing_from_framework': [],
            'potential_matches': [],
            'metadata_suggestions': {}
        }
        
        if existing_taxonomy_df is None:
            print("No existing taxonomy provided - this appears to be a new taxonomy")
            analysis['note'] = 'New taxonomy creation - no alignment comparison performed'
            return analysis
        
        print(f"Comparing {len(extracted_concepts)} extracted concepts with existing taxonomy...")
        
        # Get existing concept names
        existing_names = set()
        if hasattr(existing_taxonomy_df, 'columns'):
            # It's a DataFrame
            for col in ['Strand', 'Pillar', 'Domain', 'Skill Area', 'Skill Set', 'Skill Subset']:
                if col in existing_taxonomy_df.columns:
                    existing_names.update(existing_taxonomy_df[col].dropna().unique())
        
        analysis['existing_concept_count'] = len(existing_names)
        
        # Extract new concept names
        extracted_names = {c['MASTER_CONCEPT_NAME'] for c in extracted_concepts if c.get('MASTER_CONCEPT_NAME')}
        
        # Find overlaps and gaps
        overlapping = extracted_names.intersection(existing_names)
        missing_from_taxonomy = extracted_names - existing_names
        missing_from_framework = existing_names - extracted_names
        
        # Calculate alignment score
        if extracted_names:
            analysis['alignment_score'] = len(overlapping) / len(extracted_names)
        
        analysis['overlapping_concepts'] = list(overlapping)
        analysis['missing_from_taxonomy'] = list(missing_from_taxonomy)
        analysis['missing_from_framework'] = list(missing_from_framework)[:50]  # Limit output
        
        print(f"\n📊 Alignment Results:")
        print(f"  - Alignment score: {analysis['alignment_score']:.2%}")
        print(f"  - Overlapping concepts: {len(overlapping)}")
        print(f"  - New concepts from framework: {len(missing_from_taxonomy)}")
        print(f"  - Concepts not in framework: {len(missing_from_framework)}")
        
        return analysis
    
    def export_for_batch_mapping(self, 
                                master_concepts: List[Dict],
                                output_dir: Union[str, Path]) -> Dict:
        """
        Export master concepts in format for batch_map_skills_enhanced.py
        
        Args:
            master_concepts: List of master concept dictionaries
            output_dir: Directory to save output files
            
        Returns:
            Dictionary with paths to created files
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        print(f"\n{'='*60}")
        print(f"EXPORTING FOR BATCH MAPPING")
        print(f"{'='*60}")
        print(f"Output directory: {output_dir}")
        
        # Create concept IDs file
        concept_ids_file = output_dir / 'concept_ids.txt'
        with open(concept_ids_file, 'w') as f:
            for concept in master_concepts:
                f.write(f"{concept['MASTER_CONCEPT_ID']}\n")
        
        print(f"✅ Created concept IDs file: {concept_ids_file.name}")
        print(f"   ({len(master_concepts)} concept IDs)")
        
        # Create concepts CSV
        import csv
        concepts_csv_file = output_dir / 'concepts_for_mapping.csv'
        
        if master_concepts:
            fieldnames = list(master_concepts[0].keys())
            with open(concepts_csv_file, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(master_concepts)
        
        print(f"✅ Created concepts CSV: {concepts_csv_file.name}")
        
        # Create README
        readme_file = output_dir / 'README.md'
        readme_content = f"""# Batch Mapping Input Files

Generated: {Path(__file__).name}

## Files

- `concept_ids.txt` - List of {len(master_concepts)} concept IDs for batch processing
- `concepts_for_mapping.csv` - Full concept details in master-concepts.csv format

## Usage

Process these concepts with the batch mapping pipeline:

```bash
cd ../../analysis/scripts

python batch_map_skills_enhanced.py \\
  --concept-ids-file ../../frameworks/output/[your_output]/concept_ids.txt \\
  --content-area "English Language Arts" \\
  --checkpoint-interval 10 \\
  --output-dir ./outputs/framework_mapping
```

## Next Steps

1. Review generated concepts in `concepts_for_mapping.csv`
2. Run batch mapping pipeline (see above)
3. Review mappings and integrate into main taxonomy
"""
        
        with open(readme_file, 'w') as f:
            f.write(readme_content)
        
        print(f"✅ Created README: {readme_file.name}")
        
        return {
            'concept_ids_file': str(concept_ids_file),
            'concepts_csv_file': str(concepts_csv_file),
            'readme_file': str(readme_file),
            'concept_count': len(master_concepts),
            'output_dir': str(output_dir)
        }
    
    def generate_report(self, analysis_results: Union[Dict, List[Dict]], 
                       output_path: Optional[Path] = None) -> str:
        """
        Generate markdown report from analysis results.
        
        Args:
            analysis_results: Single or multiple analysis results
            output_path: Optional path to save report
            
        Returns:
            Markdown report string
        """
        if isinstance(analysis_results, dict):
            analysis_results = [analysis_results]
        
        lines = []
        lines.append("# Framework Analysis Report")
        lines.append("")
        lines.append(f"**Total Frameworks Analyzed:** {len(analysis_results)}")
        lines.append("")
        
        for i, result in enumerate(analysis_results, 1):
            if 'error' in result:
                lines.append(f"## Framework {i}: ERROR")
                lines.append(f"**File:** {result.get('file_path', 'unknown')}")
                lines.append(f"**Error:** {result['error']}")
                lines.append("")
                continue
            
            doc = result['document']
            lines.append(f"## Framework {i}: {doc['filename']}")
            lines.append("")
            lines.append(f"**Format:** {doc['format']}")
            lines.append(f"**Word Count:** {doc['word_count']}")
            lines.append("")
            
            # Extracted taxonomy
            lines.append("### Extracted Taxonomy")
            lines.append("```json")
            lines.append(json.dumps(result['extracted_taxonomy'], indent=2))
            lines.append("```")
            lines.append("")
            
            # Comparison if available
            if result.get('comparison'):
                lines.append("### Comparison with Our Taxonomy")
                lines.append("```json")
                lines.append(json.dumps(result['comparison'], indent=2))
                lines.append("```")
                lines.append("")
        
        # LLM usage summary
        if analysis_results and 'llm_usage' in analysis_results[0]:
            lines.append("## LLM Usage Summary")
            lines.append("")
            usage = analysis_results[0]['llm_usage']
            lines.append(f"- Total tokens: {usage['total_tokens']}")
            lines.append(f"- Estimated cost: ${usage['estimated_cost_usd']}")
            lines.append("")
        
        report = '\n'.join(lines)
        
        if output_path:
            with open(output_path, 'w') as f:
                f.write(report)
            print(f"\nReport saved to: {output_path}")
        
        return report


# Example usage
if __name__ == '__main__':
    print("=== Framework Analyzer Demo ===\n")
    print("This module requires framework documents to analyze.")
    print("\nUsage:")
    print("  from framework_analyzer import FrameworkAnalyzer")
    print("  from compatibility import TaxonomyAccess")
    print("  from llm_interface import LLMInterface")
    print("")
    print("  with TaxonomyAccess() as tax:")
    print("      llm = LLMInterface(provider='bedrock')")
    print("      analyzer = FrameworkAnalyzer(tax, llm)")
    print("      result = analyzer.analyze_framework('path/to/framework.pdf')")
    print("      report = analyzer.generate_report(result, 'framework_report.md')")

